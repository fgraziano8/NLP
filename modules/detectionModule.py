import fitz
import langdetect
import docx
from modules import configModule
from modules.objectsModule import Document
from exceptions import customExceptions


#@configModule.log_calls
def detect_language(input_text):
    try:
        langDetected = langdetect.detect(input_text)

        configModule.logging.info("Lingua rilevata: " + langDetected)
    except Exception as e:
        configModule.logging.error("Exception - Si è verificato un problema durante il processo di rilevamento della lingua - "  + str(e))
        raise customExceptions.LanguageGenericError("Errore: si è verificato un problema durante il processo di rilevamento della lingua")

    if langDetected not in configModule.config.supported_languages:
        configModule.logging.error("LanguageNotSupportedError - La lingua '" + langDetected + "' non è supportata.")
        raise customExceptions.LanguageNotSupportedError("La lingua '" + langDetected + "' non è supportata.")

    return langDetected


@configModule.log_calls
def detect_file(input_name):
    if "." not in input_name:
        configModule.logging.error("UnspecifiedExtensionError - Estensione del file non specificata.")
        raise customExceptions.UnspecifiedExtensionError("Errore: estensione del file non specificata.")

    fileType = detect_file_type(input_name)

    if fileType is None:
        configModule.logging.error("ExtensionNotSupportedError - Estensione del file non supportata.")
        raise customExceptions.ExtensionNotSupportedError("Errore: estensione del file non supportata.")
    elif fileType == "pdf":
        return detect_pdf(input_name)
    elif fileType == "docx":
        return detect_docx(input_name)
    else:
        configModule.logging.error("UnspecifiedExtensionError - Implementazione dell'estensione " + fileType + " non gestita.")
        raise customExceptions.UnspecifiedExtensionError("Errore: implementazione dell'estensione " + fileType + " non gestita.")


@configModule.log_calls
def detect_file_type(string):
    for sub in configModule.config.supported_extensions:
        if string.lower().endswith(sub):
            return sub.lower()
    return None


@configModule.log_calls
def detect_docx(input_name):
    doc_name = input_name

    try:
        if doc_name in configModule.processedDocuments:
            return doc_name

        configModule.logging.info(doc_name + " - Apertura file.")
        # Apertura del documento
        doc = docx.Document(configModule.config.document_path + doc_name)
        temp_text = []
        configModule.logging.info(doc_name + " - Aperto correttamente. Acquisizione del testo.")
        for para in doc.paragraphs:
            temp_text.append(para.text)

        text = '\n'.join(temp_text)
        text = text.replace('\n', ' ')
        configModule.logging.info(doc_name + " - Acquisito correttamente.")

        # Rileva la lingua del documento
        configModule.logging.info(doc_name + " - Rilevamento lingua.")
        lang = detect_language(text)

        newDoc = Document(doc_name, text, lang, [])
        configModule.processedDocuments[doc_name] = newDoc
        configModule.logging.info(doc_name + " - Aggiunto alla lista dei PDF elaborati")

        return doc_name

    # Gestione delle eccezioni
    except FileNotFoundError:
        configModule.logging.error("FileNotFoundError - " + doc_name + "' non trovato.")
        raise FileNotFoundError("Errore: il file '" + doc_name + "' non trovato.")
    except docx.opc.exceptions.PackageNotFoundError as a:
        configModule.logging.error("PackageNotFoundError - Impossibile leggere il file '" + doc_name + " - " + str(a))
        raise docx.opc.exceptions.PackageNotFoundError("Errore: impossibile leggere il file '" + doc_name + "'. Il documento non esiste o non è leggibile.")
    except docx.opc.exceptions.PackageNotFoundError as b:
        configModule.logging.error("PackageNotFoundError - Impossibile leggere il file '" + doc_name + " - " + str(b))
        raise docx.opc.exceptions.PackageNotFoundError("Errore: impossibile leggere il file '" + doc_name + "'. Il fileè vuoto o non contiene un documento valido.")
    except langdetect.LangDetectException as e:
        configModule.logging.error("langdetect.LangDetectException - Impossibile rilevare correttamente la lingua del file '" + doc_name + " - " + str(e))
        raise langdetect.LangDetectException("Errore: impossibile rilevare correttamente la lingua del file '" + doc_name + "'. Il testo potrebbe essere troppo breve o non avere informazioni sufficienti per una rilevazione affidabile.")
    except customExceptions.LanguageNotSupportedError:
        raise customExceptions.LanguageNotSupportedError("Errore: impossibile rilevare correttamente la lingua del file '" + doc_name + "'")
    except Exception as f:
        configModule.logging.error("Exception - Si è verificato un problema con il file '" + doc_name + " - " + str(f))
        raise Exception("Errore: si è verificato un problema nella lettura del documento " + doc_name)


@configModule.log_calls
def detect_pdf(input_name):
    pdf_name = input_name

    try:
        if pdf_name in configModule.processedDocuments:
            return pdf_name

        configModule.logging.info(pdf_name + " - Apertura file.")
        # Apertura del documento
        with fitz.Document(configModule.config.document_path + pdf_name) as doc:
            text = ""

            configModule.logging.info(pdf_name + " - Aperto correttamente. Acquisizione del testo.")
            for page in doc:
                text += page.get_text()

            text = text.replace('\n', ' ')

        configModule.logging.info(pdf_name + " - Acquisito correttamente.")

        # if not text.strip():
        #     raise RuntimeError

        # Rileva la lingua del documento
        configModule.logging.info(pdf_name + " - Rilevamento lingua.")
        lang = detect_language(text)

        newPdf = Document(pdf_name, text, lang, [])
        configModule.processedDocuments[pdf_name] = newPdf
        configModule.logging.info(pdf_name + " - Aggiunto alla lista dei PDF elaborati")

        return pdf_name

    # Gestione delle eccezioni
    except fitz.FileNotFoundError:
        configModule.logging.error("FileNotFoundError - " + pdf_name + "' non trovato.")
        print(f"Errore: il file '" + pdf_name + "' non trovato.")
        raise
    except ValueError as a:
        configModule.logging.error("ValueError - Impossibile leggere il file '" + pdf_name + " - " + str(a))
        raise ValueError("Errore: impossibile leggere il file '" + pdf_name + "'. Il documento potrebbe essere danneggiato o non può essere aperto.")
    except PermissionError as b:
        configModule.logging.error("PermissionError - Impossibile leggere il file '" + pdf_name + " - " + str(b))
        raise PermissionError("Errore: impossibile leggere il file '" + pdf_name + "'. Il documento è protetto da password.")
    except MemoryError as c:
        configModule.logging.error("MemoryError - Impossibile leggere il file '" + pdf_name + " - " + str(c))
        raise MemoryError("Errore: impossibile leggere il file '" + pdf_name + "'. Il documento è troppo grande per essere caricato in memoria.")
    # except RuntimeError as d:
    #    configModule.logging.error("RuntimeError - Impossibile estrarre il testo dal file '" + pdf_name + " - " + str(d))
    #    print(f"Errore: impossibile estrarre il testo dal file '" + pdf_name)
    #    raise
    except langdetect.LangDetectException as e:
        configModule.logging.error("langdetect.LangDetectException - Impossibile rilevare correttamente la lingua del file '" + pdf_name + " - " + str(e))
        raise langdetect.LangDetectException("Errore: impossibile rilevare correttamente la lingua del file '" + pdf_name + "'. Il testo potrebbe essere troppo breve o non avere informazioni sufficienti per una rilevazione affidabile.")
    except customExceptions.LanguageNotSupportedError:
        raise customExceptions.LanguageNotSupportedError("Errore: impossibile rilevare correttamente la lingua del file '" + pdf_name + "'")
    except Exception as f:
        configModule.logging.error("Exception - Si è verificato un problema con il file '" + pdf_name + " - " + str(f))
        raise Exception("Errore: si è verificato un problema nella lettura del documento " + pdf_name)
